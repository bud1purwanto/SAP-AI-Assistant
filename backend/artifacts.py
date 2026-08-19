"""Pembuatan berkas (Excel / CSV / Word) dari jawaban asisten.

Model tidak dapat menulis berkas biner. Sebagai gantinya model mengeluarkan
satu blok berpagar ```sap-artifact berisi JSON, lalu modul ini yang merender
berkas sungguhan dan menyimpannya di database untuk diunduh pengguna.
"""
import csv
import io
import json
import logging
import re
import uuid
from datetime import datetime, timedelta, timezone

from database import load_artifact, save_artifact

logger = logging.getLogger(__name__)

# Instruksi yang disisipkan ke system prompt.
ARTIFACT_PROMPT = """## MEMBUAT BERKAS (EXCEL / CSV / WORD)
Bila pengguna meminta hasil dalam bentuk berkas, sertakan SATU blok berikut di
akhir jawaban (selain penjelasan singkat dalam teks biasa).

**Spreadsheet** (`xlsx` atau `csv`):
```sap-artifact
{
  "type": "xlsx",
  "filename": "stok-material.xlsx",
  "sheets": [
    {"name": "Stok", "columns": ["Material", "Plant", "Qty"], "rows": [["SRRPAI", "1000", 250]]}
  ]
}
```

**Dokumen Word** (`docx`) — untuk spesifikasi, laporan, atau notulen:
```sap-artifact
{
  "type": "docx",
  "filename": "wricef-report-stok.docx",
  "title": "Functional Specification - Laporan Stok",
  "subtitle": "WRICEF ID: RPT-001",
  "meta": {"Penulis": "Tim ABAP", "Modul": "MM", "Versi": "1.0"},
  "sections": [
    {"heading": "1. Latar Belakang", "paragraphs": ["Deskripsi kebutuhan..."]},
    {"heading": "2. Ruang Lingkup", "bullets": ["Termasuk ...", "Tidak termasuk ..."]},
    {"heading": "3. Sumber Data", "table": {
       "columns": ["Tabel", "Field", "Keterangan"],
       "rows": [["MARC", "LABST", "Stok tidak terbatas"]]}}
  ]
}
```

**Dokumen WRICEF** — bila pengguna meminta dokumen WRICEF (Workflow, Report,
Interface, Conversion, Enhancement, Form), gunakan `"type": "wricef"` dengan
struktur berikut; bagian yang tidak Anda ketahui boleh dikosongkan, jangan
mengarang isinya:
```sap-artifact
{
  "type": "wricef",
  "filename": "wricef-rpt-001.docx",
  "wricef_id": "RPT-001",
  "wricef_type": "Report",
  "title": "Laporan Stok Material per Plant",
  "module": "MM",
  "author": "Tim ABAP",
  "version": "1.0",
  "business_requirement": "Uraian kebutuhan bisnis...",
  "functional_description": "Uraian cara kerja yang diharapkan...",
  "assumptions": ["Asumsi 1"],
  "selection_screen": {"columns": ["Field", "Tabel", "Wajib", "Keterangan"],
                       "rows": [["MATNR", "MARA", "Ya", "Nomor material"]]},
  "output_fields": {"columns": ["Field", "Tabel", "Keterangan"],
                    "rows": [["LABST", "MARD", "Stok tersedia"]]},
  "processing_logic": ["Langkah 1", "Langkah 2"],
  "authorization": "Objek otorisasi yang diperlukan...",
  "error_handling": [["E001", "Material tidak ditemukan"]],
  "test_scenarios": {"columns": ["No", "Skenario", "Hasil Diharapkan"],
                     "rows": [["1", "Material valid", "Data tampil"]]}
}
```

Aturan:
- Isi data dengan yang sebenarnya Anda peroleh — jangan mengarang isi tabel SAP.
- Angka ditulis sebagai angka (250), bukan teks ("250"), agar dapat dihitung di Excel.
- Jangan menampilkan blok ini bila pengguna tidak meminta berkas.
- Tetap tampilkan ringkasan isinya dalam teks agar pengguna dapat membaca
  langsung tanpa mengunduh.
"""

ARTIFACT_BLOCK_RE = re.compile(r"```sap-artifact\s*(\{.*?\})\s*```", re.DOTALL)

MAX_ROWS = 5000
MAX_COLS = 100
MAX_SECTIONS = 60
ARTIFACT_TTL = timedelta(hours=24)

SPREADSHEET_TYPES = {"xlsx", "csv"}
DOCUMENT_TYPES = {"docx", "wricef"}

CONTENT_TYPES = {
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "csv": "text/csv; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

ACCENT = "4F46E5"


def _safe_filename(name: str, default_ext: str) -> str:
    base = re.sub(r"[^A-Za-z0-9._-]+", "-", (name or "").strip()).strip("-._")
    if not base:
        base = "hasil"
    if not base.lower().endswith(f".{default_ext}"):
        base = f"{base.rsplit('.', 1)[0]}.{default_ext}"
    return base[:80]


def _normalize_cell(value):
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    return str(value)


def _as_rows(value):
    """Terima {'columns': [...], 'rows': [...]} maupun daftar baris polos."""
    if isinstance(value, dict):
        return value.get("columns") or [], value.get("rows") or []
    if isinstance(value, list):
        return [], value
    return [], []


# --- SPREADSHEET ---

def _build_xlsx(spec: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor=ACCENT)

    for sheet in spec["sheets"]:
        ws = wb.create_sheet(title=(sheet.get("name") or "Sheet1")[:31])
        columns = sheet.get("columns") or []
        rows = sheet.get("rows") or []

        if columns:
            ws.append([_normalize_cell(c) for c in columns])
            for cell in ws[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center", vertical="center")
            ws.freeze_panes = "A2"

        for row in rows:
            ws.append([_normalize_cell(c) for c in row])

        for idx in range(1, (len(columns) or (len(rows[0]) if rows else 1)) + 1):
            longest = len(str(columns[idx - 1])) if idx <= len(columns) else 0
            for row in rows[:200]:
                if idx <= len(row):
                    longest = max(longest, len(str(row[idx - 1])))
            ws.column_dimensions[get_column_letter(idx)].width = min(max(longest + 2, 10), 50)

        if columns and rows:
            ws.auto_filter.ref = f"A1:{get_column_letter(len(columns))}{len(rows) + 1}"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _build_csv(spec: dict) -> bytes:
    sheet = spec["sheets"][0]
    buf = io.StringIO()
    writer = csv.writer(buf)
    if sheet.get("columns"):
        writer.writerow(sheet["columns"])
    for row in sheet.get("rows") or []:
        writer.writerow([_normalize_cell(c) for c in row])
    # BOM agar Excel di Windows membaca UTF-8 dengan benar.
    return b"\xef\xbb\xbf" + buf.getvalue().encode("utf-8")


# --- WORD ---

def _docx_add_table(doc, columns, rows):
    """Tabel dengan baris header berwarna; melewati tabel kosong."""
    if not columns and not rows:
        return
    width = len(columns) if columns else max((len(r) for r in rows), default=0)
    if width == 0:
        return

    table = doc.add_table(rows=0, cols=width)
    table.style = "Table Grid"

    if columns:
        cells = table.add_row().cells
        for i, col in enumerate(columns[:width]):
            cells[i].text = str(col)
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i in range(width):
            cells[i].text = "" if i >= len(row) or row[i] is None else str(row[i])


def _docx_base(title: str, subtitle: str = "", meta: dict = None):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()

    heading = doc.add_heading(title or "Dokumen", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        para = doc.add_paragraph(subtitle)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.italic = True
            run.font.size = Pt(11)

    if meta:
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light List Accent 1"
        for key, value in meta.items():
            if value in (None, ""):
                continue
            cells = table.add_row().cells
            cells[0].text = str(key)
            cells[1].text = str(value)
            for para in cells[0].paragraphs:
                for run in para.runs:
                    run.bold = True
        doc.add_paragraph()

    return doc


def _build_docx(spec: dict) -> bytes:
    doc = _docx_base(spec.get("title"), spec.get("subtitle"), spec.get("meta"))

    for section in spec.get("sections") or []:
        if not isinstance(section, dict):
            continue
        if section.get("heading"):
            doc.add_heading(str(section["heading"]), level=1)
        for para in section.get("paragraphs") or []:
            doc.add_paragraph(str(para))
        for bullet in section.get("bullets") or []:
            doc.add_paragraph(str(bullet), style="List Bullet")
        for step in section.get("steps") or []:
            doc.add_paragraph(str(step), style="List Number")
        if section.get("table"):
            columns, rows = _as_rows(section["table"])
            _docx_add_table(doc, columns, rows)
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_wricef(spec: dict) -> bytes:
    """Dokumen spesifikasi WRICEF dengan urutan bagian yang lazim dipakai tim SAP."""
    wricef_id = spec.get("wricef_id") or "-"
    wricef_type = spec.get("wricef_type") or "-"

    doc = _docx_base(
        spec.get("title") or "Spesifikasi WRICEF",
        f"{wricef_type} — {wricef_id}",
        {
            "WRICEF ID": wricef_id,
            "Tipe": wricef_type,
            "Modul": spec.get("module"),
            "Penulis": spec.get("author"),
            "Versi": spec.get("version"),
            "Tanggal": spec.get("date") or datetime.now().strftime("%d %B %Y"),
        },
    )

    def text_section(heading, value):
        if not value:
            return
        doc.add_heading(heading, level=1)
        if isinstance(value, list):
            for item in value:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph(str(value))

    def table_section(heading, value):
        if not value:
            return
        columns, rows = _as_rows(value)
        if not columns and not rows:
            return
        doc.add_heading(heading, level=1)
        _docx_add_table(doc, columns, rows)
        doc.add_paragraph()

    text_section("1. Kebutuhan Bisnis", spec.get("business_requirement"))
    text_section("2. Deskripsi Fungsional", spec.get("functional_description"))
    text_section("3. Asumsi dan Batasan", spec.get("assumptions"))
    table_section("4. Layar Seleksi (Selection Screen)", spec.get("selection_screen"))
    table_section("5. Field Keluaran", spec.get("output_fields"))

    logic = spec.get("processing_logic")
    if logic:
        doc.add_heading("6. Logika Pemrosesan", level=1)
        if isinstance(logic, list):
            for step in logic:
                doc.add_paragraph(str(step), style="List Number")
        else:
            doc.add_paragraph(str(logic))

    text_section("7. Otorisasi", spec.get("authorization"))

    errors = spec.get("error_handling")
    if errors:
        columns, rows = _as_rows(errors)
        if not columns:
            columns = ["Kode", "Pesan"]
        doc.add_heading("8. Penanganan Kesalahan", level=1)
        _docx_add_table(doc, columns, rows)
        doc.add_paragraph()

    table_section("9. Skenario Pengujian", spec.get("test_scenarios"))

    for extra in spec.get("sections") or []:
        if isinstance(extra, dict) and extra.get("heading"):
            doc.add_heading(str(extra["heading"]), level=1)
            for para in extra.get("paragraphs") or []:
                doc.add_paragraph(str(para))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


BUILDERS = {
    "xlsx": _build_xlsx,
    "csv": _build_csv,
    "docx": _build_docx,
    "wricef": _build_wricef,
}


def _validate(spec: dict) -> tuple[bool, str]:
    if not isinstance(spec, dict):
        return False, "Spesifikasi berkas bukan objek JSON."

    kind = (spec.get("type") or "xlsx").lower()
    if kind not in BUILDERS:
        return False, f"Tipe berkas '{kind}' tidak didukung."

    if kind in SPREADSHEET_TYPES:
        sheets = spec.get("sheets")
        if not isinstance(sheets, list) or not sheets:
            return False, "Tidak ada sheet dalam spesifikasi berkas."
        if kind == "csv" and len(sheets) > 1:
            # Berkas CSV hanya menampung satu tabel; ambil sheet pertama.
            spec["sheets"] = sheets[:1]

        total_rows = 0
        for sheet in spec["sheets"]:
            if not isinstance(sheet, dict):
                return False, "Format sheet tidak valid."
            rows = sheet.get("rows") or []
            columns = sheet.get("columns") or []
            if not isinstance(rows, list) or not isinstance(columns, list):
                return False, "Kolom atau baris sheet tidak valid."
            if len(columns) > MAX_COLS:
                return False, f"Jumlah kolom melebihi batas {MAX_COLS}."
            total_rows += len(rows)
        if total_rows > MAX_ROWS:
            return False, f"Jumlah baris ({total_rows}) melebihi batas {MAX_ROWS}."
        return True, ""

    # Dokumen Word
    sections = spec.get("sections") or []
    if not isinstance(sections, list):
        return False, "Bagian dokumen tidak valid."
    if len(sections) > MAX_SECTIONS:
        return False, f"Jumlah bagian melebihi batas {MAX_SECTIONS}."
    if kind == "docx" and not sections and not spec.get("title"):
        return False, "Dokumen tidak memiliki judul maupun isi."
    return True, ""


def extract_and_build(reply_text: str, owner: str) -> tuple[str, list[dict]]:
    """Ambil blok artifact dari jawaban, render berkasnya, kembalikan teks bersih.

    Mengembalikan (teks_tanpa_blok, daftar_metadata_artifact).
    """
    matches = list(ARTIFACT_BLOCK_RE.finditer(reply_text or ""))
    if not matches:
        return reply_text, []

    artifacts = []
    notes = []

    for match in matches:
        try:
            spec = json.loads(match.group(1))
        except json.JSONDecodeError as e:
            logger.warning(f"Blok sap-artifact bukan JSON valid: {e}")
            notes.append("_Berkas gagal dibuat: format data dari asisten tidak valid._")
            continue

        valid, reason = _validate(spec)
        if not valid:
            logger.warning(f"Spesifikasi artifact ditolak: {reason}")
            notes.append(f"_Berkas gagal dibuat: {reason}_")
            continue

        kind = (spec.get("type") or "xlsx").lower()
        try:
            data = BUILDERS[kind](spec)
        except Exception as e:
            logger.error(f"Gagal membangun berkas {kind}: {e}")
            notes.append("_Berkas gagal dibuat karena kesalahan internal._")
            continue

        # 'wricef' hanyalah tata letak khusus dari dokumen Word.
        extension = "docx" if kind in DOCUMENT_TYPES else kind
        artifact_id = uuid.uuid4().hex
        filename = _safe_filename(spec.get("filename"), extension)

        stored = save_artifact(
            artifact_id=artifact_id,
            owner=owner,
            filename=filename,
            content_type=CONTENT_TYPES[extension],
            kind=extension,
            data=data,
            expires_at=datetime.now(timezone.utc) + ARTIFACT_TTL,
        )
        if not stored:
            notes.append("_Berkas gagal disimpan, silakan coba lagi._")
            continue

        artifacts.append({
            "artifact_id": artifact_id,
            "filename": filename,
            "type": extension,
            "size": len(data),
        })

    cleaned = ARTIFACT_BLOCK_RE.sub("", reply_text).strip()
    if notes:
        cleaned = f"{cleaned}\n\n" + "\n".join(notes)
    return cleaned, artifacts


def get_artifact(artifact_id: str, owner: str):
    """Ambil berkas milik user tertentu; None bila tidak ada/kedaluwarsa/bukan miliknya."""
    return load_artifact(artifact_id, owner)

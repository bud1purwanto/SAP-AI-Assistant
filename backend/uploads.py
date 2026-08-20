"""Lampiran percakapan: berkas dan gambar yang dikirim pengguna sebagai konteks.

Berkas disimpan di PostgreSQL (sama seperti berkas hasil generate) agar tetap
dapat diakses oleh semua worker uvicorn dan bertahan melewati restart.

Teks diekstraksi sekali saat unggah, bukan pada setiap giliran percakapan:
ekstraksi PDF/Word/Excel relatif mahal dan hasilnya tidak berubah.
"""
import csv
import io
import json
import logging
import re
import uuid
from datetime import datetime, timedelta, timezone

from config import settings
from database import save_upload

logger = logging.getLogger(__name__)

UPLOAD_TTL = timedelta(days=30)

# Batas per berkas dan per pesan; melindungi database sekaligus context window model.
MAX_FILE_BYTES = 10 * 1024 * 1024      # 10 MB
MAX_IMAGE_BYTES = 5 * 1024 * 1024      # 5 MB
MAX_TEXT_CHARS = 40_000                # ~10k token per berkas
MAX_ATTACHMENTS_PER_MESSAGE = 5

IMAGE_TYPES = {
    "image/png": "png",
    "image/jpeg": "jpg",
    "image/webp": "webp",
    "image/gif": "gif",
}

DOCUMENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "text/csv": "csv",
    "text/plain": "txt",
    "text/markdown": "md",
    "application/json": "json",
}

# Sebagian browser mengirim MIME kosong atau generik; ekstensi dipakai sebagai cadangan.
EXTENSION_TYPES = {
    "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
    "webp": "image/webp", "gif": "image/gif", "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "csv": "text/csv", "txt": "text/plain", "md": "text/markdown",
    "json": "application/json",
}


class UploadRejected(Exception):
    """Berkas ditolak beserta alasan yang dapat ditampilkan ke pengguna."""


def safe_filename(name: str) -> str:
    base = re.sub(r"[^A-Za-z0-9._-]+", "-", (name or "").strip()).strip("-._")
    return (base or "berkas")[:120]


def resolve_content_type(filename: str, declared: str) -> str:
    """Tentukan tipe berkas dari MIME yang dikirim browser, dengan cadangan ekstensi."""
    declared = (declared or "").split(";")[0].strip().lower()
    if declared in IMAGE_TYPES or declared in DOCUMENT_TYPES:
        return declared

    extension = (filename or "").rsplit(".", 1)[-1].lower() if "." in (filename or "") else ""
    return EXTENSION_TYPES.get(extension, declared or "application/octet-stream")


def _truncate(text: str) -> tuple[str, bool]:
    text = (text or "").strip()
    if len(text) <= MAX_TEXT_CHARS:
        return text, False
    return text[:MAX_TEXT_CHARS], True


# --- Ekstraksi teks ---

def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            content = page.extract_text() or ""
        except Exception as e:  # halaman rusak tidak boleh menggagalkan seluruh berkas
            logger.warning(f"Gagal membaca halaman {index} PDF: {e}")
            continue
        if content.strip():
            pages.append(f"[Halaman {index}]\n{content.strip()}")
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    sheets = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            if row is None:
                continue
            cells = ["" if c is None else str(c) for c in row]
            if any(c.strip() for c in cells):
                rows.append(" | ".join(cells))
            if len(rows) >= 500:  # cukup untuk konteks; berkas besar tidak dimuat penuh
                rows.append("… (baris berikutnya tidak disertakan)")
                break
        if rows:
            sheets.append(f"[Sheet: {ws.title}]\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets)


def _extract_csv(data: bytes) -> str:
    text = data.decode("utf-8-sig", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = []
    for index, row in enumerate(reader):
        rows.append(" | ".join(row))
        if index >= 500:
            rows.append("… (baris berikutnya tidak disertakan)")
            break
    return "\n".join(rows)


def _extract_json(data: bytes) -> str:
    raw = data.decode("utf-8", errors="replace")
    try:
        return json.dumps(json.loads(raw), indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        return raw


EXTRACTORS = {
    "application/pdf": _extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": _extract_xlsx,
    "text/csv": _extract_csv,
    "application/json": _extract_json,
    "text/plain": lambda data: data.decode("utf-8", errors="replace"),
    "text/markdown": lambda data: data.decode("utf-8", errors="replace"),
}


def extract_text(content_type: str, data: bytes) -> tuple[str, str]:
    """Kembalikan (teks, catatan). Catatan berisi alasan bila teks kosong/terpotong."""
    extractor = EXTRACTORS.get(content_type)
    if not extractor:
        return "", ""

    try:
        text, truncated = _truncate(extractor(data))
    except Exception as e:
        logger.error(f"Ekstraksi teks gagal untuk {content_type}: {e}")
        return "", "Isi berkas tidak dapat dibaca otomatis."

    if not text:
        return "", "Berkas tidak memuat teks yang dapat dibaca (kemungkinan hasil pindaian)."
    if truncated:
        return text, f"Hanya {MAX_TEXT_CHARS:,} karakter pertama yang disertakan."
    return text, ""


def store_upload(owner: str, filename: str, declared_type: str, data: bytes,
                 session_id: str = None) -> dict:
    """Validasi, ekstraksi, dan simpan satu berkas. Melempar UploadRejected bila ditolak."""
    if not data:
        raise UploadRejected("Berkas kosong.")

    content_type = resolve_content_type(filename, declared_type)
    is_image = content_type in IMAGE_TYPES

    if not is_image and content_type not in DOCUMENT_TYPES:
        raise UploadRejected(
            "Jenis berkas belum didukung. Yang bisa dilampirkan: gambar (PNG, JPG, WEBP, GIF), "
            "PDF, Word, Excel, CSV, teks, Markdown, dan JSON."
        )

    limit = MAX_IMAGE_BYTES if is_image else MAX_FILE_BYTES
    if len(data) > limit:
        raise UploadRejected(f"Ukuran berkas melebihi batas {limit // (1024 * 1024)} MB.")

    text, note = ("", "") if is_image else extract_text(content_type, data)

    upload_id = uuid.uuid4().hex
    stored = save_upload(
        upload_id=upload_id,
        owner=owner,
        session_id=session_id,
        filename=safe_filename(filename),
        content_type=content_type,
        kind="image" if is_image else "document",
        data=data,
        extracted_text=text,
        expires_at=datetime.now(timezone.utc) + UPLOAD_TTL,
    )
    if not stored:
        raise UploadRejected("Berkas gagal disimpan, silakan coba lagi.")

    return {
        "upload_id": upload_id,
        "filename": safe_filename(filename),
        "content_type": content_type,
        "kind": "image" if is_image else "document",
        "size": len(data),
        "has_text": bool(text),
        "note": note,
    }


def build_context_blocks(uploads: list[dict]) -> tuple[str, list[dict]]:
    """Ubah lampiran menjadi (blok teks untuk prompt, daftar gambar untuk model vision)."""
    text_parts = []
    images = []

    for item in uploads:
        if item["kind"] == "image":
            import base64

            images.append({
                "filename": item["filename"],
                "data_url": (
                    f"data:{item['content_type']};base64,"
                    + base64.b64encode(item["data"]).decode("ascii")
                ),
            })
            continue

        body = (item.get("extracted_text") or "").strip()
        if body:
            text_parts.append(f"--- ISI BERKAS: {item['filename']} ---\n{body}\n--- AKHIR BERKAS ---")
        else:
            text_parts.append(
                f"--- BERKAS: {item['filename']} (tidak ada teks yang dapat dibaca) ---"
            )

    return "\n\n".join(text_parts), images

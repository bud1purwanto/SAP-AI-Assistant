"""Lampiran percakapan: unggah, ekstraksi teks, dan pembatasan akses."""
import io
import json

import pytest


def _png_bytes():
    """PNG 1x1 minimal yang valid."""
    import base64

    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
    )


def _xlsx_bytes():
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Stok"
    ws.append(["Material", "Qty"])
    ws.append(["SRRPAI", 250])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _docx_bytes():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Prosedur pengadaan barang tahun 2026.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Langkah"
    table.rows[0].cells[1].text = "Buat PO"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _upload(client, auth, name, data, content_type):
    return client.post(
        "/api/uploads",
        files={"file": (name, data, content_type)},
        headers=auth,
    )


def test_text_file_is_extracted(client, make_user):
    auth = make_user("up_txt")
    res = _upload(client, auth, "catatan.txt", "Stok material SRRPAI adalah 250 PC.".encode(), "text/plain")
    assert res.status_code == 200, res.text
    body = res.json()
    assert body["kind"] == "document" and body["has_text"] is True


def test_spreadsheet_content_becomes_context(client, make_user, db):
    auth = make_user("up_xlsx")
    res = _upload(
        client, auth, "stok.xlsx", _xlsx_bytes(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    assert res.status_code == 200, res.text

    stored = db.load_uploads([res.json()["upload_id"]], "up_xlsx")[0]
    assert "SRRPAI" in stored["extracted_text"] and "250" in stored["extracted_text"]


def test_word_document_content_becomes_context(client, make_user, db):
    auth = make_user("up_docx")
    res = _upload(
        client, auth, "sop.docx", _docx_bytes(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    stored = db.load_uploads([res.json()["upload_id"]], "up_docx")[0]
    assert "pengadaan barang" in stored["extracted_text"]
    assert "Buat PO" in stored["extracted_text"]  # isi tabel ikut terbaca


def test_image_is_stored_without_text_extraction(client, make_user):
    auth = make_user("up_img")
    res = _upload(client, auth, "layar.png", _png_bytes(), "image/png")
    assert res.status_code == 200, res.text
    assert res.json()["kind"] == "image" and res.json()["has_text"] is False


def test_unsupported_type_rejected(client, make_user):
    auth = make_user("up_bad")
    res = _upload(client, auth, "installer.exe", b"MZ\x90\x00", "application/x-msdownload")
    assert res.status_code == 400
    assert "belum didukung" in res.json()["detail"]


def test_oversized_file_rejected(client, make_user):
    from uploads import MAX_FILE_BYTES

    auth = make_user("up_big")
    res = _upload(client, auth, "besar.txt", b"x" * (MAX_FILE_BYTES + 1), "text/plain")
    assert res.status_code == 400 and "melebihi batas" in res.json()["detail"]


def test_attachment_is_private_to_its_owner(client, make_user):
    owner = make_user("att_owner")
    other = make_user("att_other")
    upload_id = _upload(client, owner, "rahasia.txt", b"gaji karyawan", "text/plain").json()["upload_id"]

    assert client.get(f"/api/uploads/{upload_id}", headers=owner).status_code == 200
    # User lain: 404, bukan 403 — keberadaan lampiran orang lain tidak dibocorkan.
    assert client.get(f"/api/uploads/{upload_id}", headers=other).status_code == 404
    # Tanpa token: 401, karena kredensialnya yang tidak ada.
    assert client.get(f"/api/uploads/{upload_id}").status_code == 401


def test_load_uploads_ignores_other_users_ids(client, make_user, db):
    owner = make_user("load_owner")
    make_user("load_other")
    upload_id = _upload(client, owner, "a.txt", b"isi", "text/plain").json()["upload_id"]

    assert db.load_uploads([upload_id], "load_other") == []


def test_extension_used_when_browser_sends_no_mime(client, make_user):
    """Sebagian browser mengirim content-type kosong."""
    auth = make_user("up_mime")
    res = _upload(client, auth, "data.csv", b"a,b\n1,2\n", "application/octet-stream")
    assert res.status_code == 200, res.text
    assert res.json()["content_type"] == "text/csv"


def test_context_blocks_split_text_and_images(db, client, make_user):
    from uploads import build_context_blocks

    auth = make_user("ctx_user")
    doc_id = _upload(client, auth, "nota.txt", b"Isi dokumen penting", "text/plain").json()["upload_id"]
    img_id = _upload(client, auth, "foto.png", _png_bytes(), "image/png").json()["upload_id"]

    loaded = db.load_uploads([doc_id, img_id], "ctx_user")
    text_block, images = build_context_blocks(loaded)

    assert "Isi dokumen penting" in text_block and "nota.txt" in text_block
    assert len(images) == 1 and images[0]["data_url"].startswith("data:image/png;base64,")


# --- Lampiran sebagai konteks model ---

def _captured_agent(monkeypatch, fail_on_images=False):
    """Jalankan agen dengan LLM tiruan dan tangkap pesan yang dikirim ke model."""
    import agent

    seen = {"calls": []}

    class FakeTool:
        name = "read_table"
        description = "d"
        inputSchema = {"type": "object", "properties": {}}

    class FakeLLM:
        def bind_tools(self, *args, **kwargs):
            return self

        async def ainvoke(self, messages):
            seen["calls"].append(messages[-1].content)
            has_image = isinstance(messages[-1].content, list)
            if fail_on_images and has_image:
                raise RuntimeError("400: this model does not support image_url content")

            class Response:
                content = "Jawaban."
                tool_calls = []

            return Response()

    async def fake_tools(server_filter="all"):
        return [{"server": "sap", "tool": FakeTool()}]

    monkeypatch.setattr(agent.mcp_manager, "get_all_tools", fake_tools)
    monkeypatch.setattr(agent, "ChatOpenAI", lambda **kwargs: FakeLLM())
    return seen


def test_document_text_reaches_the_model(client, make_user, db, monkeypatch):
    import asyncio

    import agent
    from models import ChatRequest

    auth = make_user("agen_doc")
    upload_id = _upload(
        client, auth, "sop.txt", b"Batas minimum stok gudang adalah 100 PC.", "text/plain"
    ).json()["upload_id"]

    seen = _captured_agent(monkeypatch)
    asyncio.run(
        agent.process_chat(
            ChatRequest(message="Berapa batas minimum stok?", attachment_ids=[upload_id]),
            "user",
            "",
            username="agen_doc",
        )
    )

    sent = seen["calls"][0]
    assert "Batas minimum stok gudang adalah 100 PC." in sent
    assert "sop.txt" in sent


def test_image_is_sent_as_a_vision_part(client, make_user, monkeypatch):
    import asyncio

    import agent
    from models import ChatRequest

    auth = make_user("agen_img")
    upload_id = _upload(client, auth, "layar.png", _png_bytes(), "image/png").json()["upload_id"]

    seen = _captured_agent(monkeypatch)
    asyncio.run(
        agent.process_chat(
            ChatRequest(message="Apa isi gambar ini?", attachment_ids=[upload_id]),
            "user",
            "",
            username="agen_img",
        )
    )

    sent = seen["calls"][0]
    assert isinstance(sent, list), "gambar harus dikirim sebagai bagian multimodal"
    kinds = [part["type"] for part in sent]
    assert "image_url" in kinds and "text" in kinds
    image_part = [p for p in sent if p["type"] == "image_url"][0]
    assert image_part["image_url"]["url"].startswith("data:image/png;base64,")


def test_falls_back_to_text_when_model_rejects_images(client, make_user, monkeypatch):
    """Model tanpa dukungan gambar tidak boleh menggagalkan seluruh percakapan."""
    import asyncio

    import agent
    from models import ChatRequest

    auth = make_user("agen_fallback")
    upload_id = _upload(client, auth, "layar.png", _png_bytes(), "image/png").json()["upload_id"]

    seen = _captured_agent(monkeypatch, fail_on_images=True)
    result = asyncio.run(
        agent.process_chat(
            ChatRequest(message="Apa isi gambar ini?", attachment_ids=[upload_id]),
            "user",
            "",
            username="agen_fallback",
        )
    )

    assert len(seen["calls"]) == 2, "seharusnya dicoba ulang tanpa gambar"
    assert isinstance(seen["calls"][1], str), "percobaan kedua harus teks saja"
    assert "tidak dapat diproses" in seen["calls"][1]
    assert result.reply  # pengguna tetap menerima jawaban


def test_attachments_of_another_user_are_ignored(client, make_user, monkeypatch):
    """ID lampiran datang dari klien dan tidak boleh dipercaya begitu saja."""
    import asyncio

    import agent
    from models import ChatRequest

    owner = make_user("agen_owner")
    make_user("agen_penyusup")
    upload_id = _upload(client, owner, "rahasia.txt", b"RAHASIA PERUSAHAAN", "text/plain").json()["upload_id"]

    seen = _captured_agent(monkeypatch)
    asyncio.run(
        agent.process_chat(
            ChatRequest(message="Apa isinya?", attachment_ids=[upload_id]),
            "user",
            "",
            username="agen_penyusup",
        )
    )

    assert "RAHASIA PERUSAHAAN" not in str(seen["calls"][0])

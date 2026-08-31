"""Pembuatan berkas Excel / CSV / Word dan aturan pengunduhannya."""
import io
import json

import pytest
from docx import Document
from openpyxl import load_workbook

XLSX_SPEC = {
    "type": "xlsx",
    "filename": "stok material.xlsx",
    "sheets": [{"name": "Stok", "columns": ["Material", "Qty"], "rows": [["SRRPAI", 250]]}],
}

WRICEF_SPEC = {
    "type": "wricef",
    "filename": "wricef-rpt-001.docx",
    "wricef_id": "RPT-001",
    "wricef_type": "Report",
    "title": "Laporan Stok",
    "module": "MM",
    "business_requirement": "Tim gudang perlu laporan stok harian.",
    "processing_logic": ["Baca MARC", "Tampilkan ALV"],
    "test_scenarios": {"columns": ["No", "Skenario"], "rows": [["1", "Material valid"]]},
}


def _block(spec):
    return "Ringkasan.\n\n```sap-artifact\n" + json.dumps(spec) + "\n```"


@pytest.fixture
def artifacts(db):
    import artifacts as module

    return module


def test_xlsx_has_real_values_and_numeric_cells(artifacts):
    text, built = artifacts.extract_and_build(_block(XLSX_SPEC), owner="andi")

    assert "```sap-artifact" not in text and "Ringkasan." in text
    assert built[0]["filename"] == "stok-material.xlsx"

    data = artifacts.get_artifact(built[0]["artifact_id"], "andi")["data"]
    sheet = load_workbook(io.BytesIO(data))["Stok"]
    assert [c.value for c in sheet[1]] == ["Material", "Qty"]
    # Angka harus tersimpan sebagai angka agar dapat dihitung di Excel.
    assert sheet["B2"].value == 250 and isinstance(sheet["B2"].value, int)


def test_csv_carries_bom_for_excel(artifacts):
    spec = {"type": "csv", "filename": "data", "sheets": [{"columns": ["A"], "rows": [[1]]}]}
    _, built = artifacts.extract_and_build(_block(spec), owner="andi")
    assert built[0]["filename"] == "data.csv"
    assert artifacts.get_artifact(built[0]["artifact_id"], "andi")["data"].startswith(b"\xef\xbb\xbf")


def test_wricef_document_has_expected_sections(artifacts):
    _, built = artifacts.extract_and_build(_block(WRICEF_SPEC), owner="andi")
    assert built[0]["type"] == "docx"

    item = artifacts.get_artifact(built[0]["artifact_id"], "andi")
    assert item["data"][:2] == b"PK"

    doc = Document(io.BytesIO(item["data"]))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith(("Heading", "Title"))]
    for expected in ["Kebutuhan Bisnis", "Logika Pemrosesan", "Skenario Pengujian"]:
        assert any(expected in h for h in headings), (expected, headings)

    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    assert "RPT-001" in cells


def test_artifact_is_private_to_its_owner(artifacts):
    _, built = artifacts.extract_and_build(_block(XLSX_SPEC), owner="andi")
    assert artifacts.get_artifact(built[0]["artifact_id"], "siti") is None


def test_artifact_survives_a_different_process(artifacts, db):
    """Penyimpanan di memori gagal saat uvicorn berjalan dengan banyak worker."""
    _, built = artifacts.extract_and_build(_block(XLSX_SPEC), owner="andi")
    artifact_id = built[0]["artifact_id"]

    # Baca langsung dari database, bukan dari state proses.
    assert db.load_artifact(artifact_id, "andi") is not None


def test_malformed_spec_keeps_the_answer(artifacts):
    text, built = artifacts.extract_and_build("Jawaban penting.\n```sap-artifact\n{rusak}\n```", owner="andi")
    assert built == []
    assert "Jawaban penting." in text and "gagal dibuat" in text


def test_oversized_spec_rejected(artifacts):
    spec = {"type": "xlsx", "sheets": [{"columns": ["A"], "rows": [[1]] * 6000}]}
    _, built = artifacts.extract_and_build(_block(spec), owner="andi")
    assert built == []


def test_empty_document_rejected(artifacts):
    _, built = artifacts.extract_and_build('x\n```sap-artifact\n{"type": "docx"}\n```', owner="andi")
    assert built == []


def test_per_user_quota_drops_oldest(artifacts, db):
    from config import settings

    owner = "kuota_user"
    for i in range(settings.artifact_max_per_user + 5):
        spec = dict(XLSX_SPEC, filename=f"berkas-{i}.xlsx")
        artifacts.extract_and_build(_block(spec), owner=owner)

    assert db.count_user_artifacts(owner) <= settings.artifact_max_per_user
